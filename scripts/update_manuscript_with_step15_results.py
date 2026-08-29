from pathlib import Path
from docx import Document


PROJECT = Path(r".")
SRC = PROJECT / "results" / "11_manuscript_preparation" / "14A_full_manuscript_draft" / "14A_full_manuscript_draft_v2_with_figures.docx"
DST = PROJECT / "results" / "11_manuscript_preparation" / "14A_full_manuscript_draft" / "14A_full_manuscript_draft_v4_strengthened_with_figures.docx"


def replace_paragraph_text(paragraph, text):
    paragraph.text = text


def find_first(doc, predicate):
    for paragraph in doc.paragraphs:
        if predicate(paragraph.text):
            return paragraph
    raise RuntimeError("Target paragraph was not found")


def insert_before(anchor, text, style=None):
    paragraph = anchor.insert_paragraph_before(text)
    if style:
        paragraph.style = style
    return paragraph


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


doc = Document(str(SRC))

for section in doc.sections:
    for paragraph in section.header.paragraphs:
        if "Initial manuscript draft" in paragraph.text:
            replace_paragraph_text(
                paragraph,
                paragraph.text.replace(
                    "Initial manuscript draft", "Strengthened manuscript draft"
                ),
            )

for paragraph in doc.paragraphs:
    if "ACTOMYOSIN_ROBUST_INTEGRIN_NON_SPECIFIC ." in paragraph.text:
        replace_paragraph_text(
            paragraph,
            paragraph.text.replace(
                "ACTOMYOSIN_ROBUST_INTEGRIN_NON_SPECIFIC .",
                "ACTOMYOSIN_ROBUST_INTEGRIN_NON_SPECIFIC.",
            ),
        )

for paragraph in doc.paragraphs:
    if paragraph.text.startswith("The appropriate interpretation is therefore an evidence-bounded association hypothesis"):
        replace_paragraph_text(
            paragraph,
            "The appropriate interpretation is therefore an evidence-bounded association hypothesis rather than a demonstrated causal mechanosensitivity mechanism. The final project evidence grade remained CAUTION; the integrated interpretation was that actomyosin support was robust whereas integrin specificity was not established.",
        )

replace_paragraph_text(
    find_first(doc, lambda x: x.startswith("Manuscript status:")),
    "Manuscript status: evidence-grounded strengthened draft v4; current evidence grade CAUTION.",
)
replace_paragraph_text(
    find_first(doc, lambda x: x.startswith("We performed a provenance-aware reanalysis")),
    "We performed a provenance-aware reanalysis of public human transcriptomic resources comprising discovery analysis in GSE173252, paired sample-level validation in 12 PRJNA607098 sample-state units, independent external triangulation in five GSE130973 subjects, and a prespecified mechanochemical perturbation analysis of GSE300230. Frozen candidate genes and predefined modules were evaluated with sample/donor-level evidence boundaries.",
)
replace_paragraph_text(
    find_first(doc, lambda x: x.startswith("The core module and leave-one-sample-out gates passed")),
    "The core module and leave-one-sample-out gates passed in the primary validation. Candidate support was 10/15 overall and 5/6 for the integrin/focal-adhesion branch, but competition robustness failed and PIEZO2 was not supported in the primary source. In GSE130973, integrin support was 5/6, actomyosin candidate support was 1/4, and PIEZO2 was directionally positive. In GSE300230, mechanical tension increased the actomyosin module in both fibroblast cell lines, whereas integrin support was directional but less uniform; cell-cycle activation remained a competing explanation.",
)
replace_paragraph_text(
    find_first(doc, lambda x: x.startswith("The data support a bounded fibroblast-associated")),
    "The data support a bounded fibroblast-associated ECM/integrin–actomyosin transcriptional program and make a tension-associated actomyosin response biologically plausible in cultured dermal fibroblasts. They do not establish a universal single-gene mechanism, fascia specificity, or causal mechanosensitivity. Overall evidence grade: CAUTION.",
)

replace_paragraph_text(
    find_first(doc, lambda x: x.startswith("We conducted a provenance-aware, multi-stage reanalysis")),
    "We conducted a provenance-aware, multi-stage reanalysis of publicly available human transcriptomic resources to evaluate whether fascia-like or myofibroblast-enriched fibroblast states are associated with a multicomponent extracellular-matrix/integrin–cytoskeletal mechanotransduction program. The workflow comprised discovery analysis in GSE173252, paired sample-level validation using the PRJNA607098-derived atlas representation, independent external triangulation using GSE130973, and a prespecified mechanochemical perturbation analysis of GSE300230. The analysis was designed to separate module-level association from candidate-gene-level direction and from causal mechanosensitivity.",
)

gse130_anchor = find_first(doc, lambda x: x == "GSE130973 external triangulation")
evidence_anchor = find_first(doc, lambda x: x == "Evidence synthesis and interpretation rules")
insert_before(
    evidence_anchor,
    "GSE300230 mechanochemical perturbation triangulation",
    "Heading 2",
)
insert_before(
    evidence_anchor,
    "GSE300230 provided a processed 16-sample raw-count matrix with a complete 2 × 2 mechanical-tension-by-TGF-beta design in each of two primary dermal fibroblast cell lines. GM08401 represented an older donor background and GM09503 represented a young donor background, with age completely confounded with cell line. We therefore analyzed the two cell lines separately, included batch as a blocking factor, and used tension versus relaxed conditions without exogenous TGF-beta as the primary mechanical contrast. TGF-beta-present mechanical contrasts and the mechanical-by-TGF-beta interaction were secondary analyses. No population-level age effect was estimated.",
)
insert_before(
    evidence_anchor,
    "The prespecified module and candidate panel, alias rules, filtering rule, and multiplicity families were frozen before condition-specific expression testing. Module-level inference used limma-voom and CAMERA with correction across the nine frozen modules within each cell line and contrast. Candidate-level inference used the unchanged 15-gene panel with correction within that family. All results were interpreted as perturbation plausibility in cultured dermal fibroblasts rather than as fascia-specific or causal evidence.",
)

replace_paragraph_text(
    find_first(doc, lambda x: x.startswith("The final evidence synthesis integrated provenance")),
    "The final evidence synthesis integrated provenance, sample structure, module-level results, candidate-level direction, competition robustness, cross-source concordance, and the bounded GSE300230 perturbation analysis. The interpretation classes were assigned before converting the results into manuscript claims. The current evidence grade was retained as CAUTION because competition robustness was incomplete, PIEZO2 was source-dependent, actomyosin candidate genes were discordant across sources, cell-cycle competition remained in the perturbation source, and direct F7/F8 equivalence remained unresolved.",
)
replace_paragraph_text(
    find_first(doc, lambda x: x.startswith("The revised primary hypothesis was therefore stated")),
    "The revised primary hypothesis was therefore stated at the module level: fascia-like or myofibroblast-enriched fibroblast states may exhibit a context-dependent ECM/integrin–actomyosin transcriptional program, and mechanical tension may induce an actomyosin-associated response in cultured human dermal fibroblasts. Individual genes, channel-specific components, age effects, cell-intrinsic effects, and fascia-specific causality remain unresolved. Functional mechanical perturbation, donor-level replication, and orthogonal readouts are required before causal mechanosensitivity can be claimed.",
)

results_anchor = find_first(doc, lambda x: x == "Integrated evidence boundary")
insert_before(results_anchor, "Mechanochemical perturbation triangulation in GSE300230", "Heading 2")
insert_before(
    results_anchor,
    "The prespecified GSE300230 analysis provided a perturbation test complementary to the observational single-cell analyses. In the primary contrast, mechanical tension versus relaxed conditions without exogenous TGF-beta increased the actomyosin/Rho module in both cell lines, with module-family q values of 5.07 × 10−9 for GM08401 and 1.20 × 10−5 for GM09503. The integrin/focal-adhesion module was also positive in both lines, although its module-family q values were 2.67 × 10−4 and 0.115, respectively. The frozen candidate panel contained eight of fifteen genes with positive effects in both lines, and no candidate reached q ≤ 0.05 in both lines.",
)
insert_before(
    results_anchor,
    "The perturbation result was stable to the prespecified score-sensitivity analysis: both core module scores retained positive direction under z-mean, z-median, rank-mean and centered-mean summaries, and leave-one-gene-out analyses preserved positive direction for both core modules in both cell lines. However, the cell-cycle module was also increased in both lines for the primary contrast, with q values of 0.00225 and 5.12 × 10−12. Thus, the data support actomyosin perturbation plausibility but do not separate it from a coupled proliferation-associated response. The cell-intrinsic versus composition-associated explanation was not estimable from the gene-by-sample cultured-fibroblast matrix.",
)
replace_paragraph_text(
    find_first(doc, lambda x: x.startswith("The integrated interpretation was therefore restricted")),
    "The integrated interpretation was therefore restricted to a module-level fibroblast-associated ECM/integrin–cytoskeletal program. The final evidence grade remained CAUTION. The available data do not establish causal mechanosensitivity, direct F7/F8 replication, universal PIEZO2 involvement, or specificity from generic fibroblast remodeling and inflammatory states. The new perturbation evidence narrows rather than overturns this interpretation: it supports an actomyosin-associated response to mechanical tension in two cultured dermal fibroblast cell lines, while leaving proliferation competition, donor generalization, age inference and fascia specificity unresolved.",
)

replace_paragraph_text(
    find_first(doc, lambda x: x.startswith("This study provides evidence for a fibroblast-associated")),
    "This study provides evidence for a fibroblast-associated, multicomponent ECM/integrin–cytoskeletal program in selected human single-cell sources and a tension-associated actomyosin response in cultured dermal fibroblasts. The most stable signal was observed at the module level rather than as a universally conserved set of individual genes. Paired sample-level validation supported the core program and leave-one-sample-out robustness, whereas the competition-robustness gate failed. Independent GSE130973 triangulation provided partial support for the integrin branch but revealed actomyosin candidate-gene discordance and a source-dependent PIEZO2 direction.",
)
alternative_anchor = find_first(doc, lambda x: x == "Alternative explanations and specificity limits")
insert_before(alternative_anchor, "Perturbation evidence strengthens actomyosin plausibility but exposes a proliferation competitor", "Heading 2")
insert_before(
    alternative_anchor,
    "GSE300230 adds a distinct type of evidence because mechanical tension was experimentally contrasted within two cultured primary dermal fibroblast cell lines. The actomyosin/Rho module showed a strong, directionally concordant tension response in both lines, and this direction was stable across scoring summaries and leave-one-gene-out analyses. This supports prioritizing contractility, traction and Rho/ROCK-dependent readouts for functional follow-up.",
)
insert_before(
    alternative_anchor,
    "The same perturbation contrast also increased the cell-cycle module in both lines. The expression-matched negative-control analysis did not remove the actomyosin score effect, but it cannot establish that the effect is independent of proliferation because cell-cycle and contractility programs may be coupled in activated fibroblasts. The appropriate claim is therefore perturbation plausibility, not pathway-specific causal activation.",
)
replace_paragraph_text(
    find_first(doc, lambda x: x.startswith("The current design cannot distinguish these explanations")),
    "The current design cannot distinguish these explanations by transcriptomic association alone. GSE300230 also cannot separate cell-intrinsic regulation from cell-composition changes because it is a gene-by-sample cultured-fibroblast matrix without cell-level composition measurements. The next experiment must therefore manipulate mechanical context while measuring viability, proliferation, ECM/TGF/inflammation state, and an orthogonal functional outcome in the same donor-matched system.",
)

replace_paragraph_text(
    find_first(doc, lambda x: x.startswith("Several limitations constrain the interpretation")),
    "Several limitations constrain the interpretation. First, the analysis relied partly on processed public objects and remote gene-panel extraction rather than a harmonized raw-data reprocessing pipeline. Second, the atlas F7 state and paper F8 state were not proven equivalent. Third, GSE130973 was a marker-defined, aging-skin fibroblast-state audit and therefore provides external triangulation rather than direct replication. Fourth, GSE175817 could not be used for donor-level validation because donor-column mapping remained unresolved. Fifth, GSE300230 contained only two cell lines, with age completely confounded with cell line, and therefore cannot support an independent age effect or donor-population generalization. Sixth, the GSE300230 cell-cycle response remained a competing explanation and cell-intrinsic versus composition-associated effects were not estimable. Finally, all current evidence remains bounded and cannot establish fascia-specific or pain-relevant causality.",
)

conclusion_heading = find_first(doc, lambda x: x.strip() == "Conclusion")
conclusion_bodies = [
    p for p in doc.paragraphs
    if p.text.startswith("The study supports a bounded hypothesis") or p.text.startswith("This study identifies a bounded")
]
if not conclusion_bodies:
    raise RuntimeError("Conclusion body was not found")
replace_paragraph_text(
    conclusion_bodies[0],
    "This study supports a bounded hypothesis that fascia-like fibroblast states are associated with a context-dependent ECM/integrin–actomyosin transcriptional program. GSE300230 adds perturbation plausibility for a tension-associated actomyosin response in cultured dermal fibroblasts, but integrin specificity, mechanosensor involvement, age effects, cell-intrinsic regulation, fascia specificity and pain-relevant causality remain unresolved. The decisive next step is donor-replicated functional mechanical perturbation with proliferation and state controls, prespecified integrin/actomyosin perturbation, and orthogonal contractility or traction readouts.",
)
for body in conclusion_bodies[1:]:
    remove_paragraph(body)

replace_paragraph_text(
    find_first(doc, lambda x: x.startswith("The analysis reuses public resources identified by the GEO/NCBI accessions")),
    "The analysis reuses public resources identified by the GEO/NCBI accessions GSE173252, GSE273293, PRJNA607098, GSE130973, GSE300230, and the screened independent sources described in the provenance audit. The project-local analysis scripts, configuration files, audit tables, and result files are stored under `.`. A public repository URL for the complete code release has not yet been assigned and must be added before submission.",
)

doc.save(str(DST))
print(f"WROTE {DST}")
print(f"PARAGRAPHS {len(doc.paragraphs)}")
print(f"INLINE_SHAPES {len(doc.inline_shapes)}")
