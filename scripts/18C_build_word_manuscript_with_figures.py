from pathlib import Path
import os
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path(r".")
INPUT_MD = Path(os.environ.get(
    "MANUSCRIPT_INPUT_MD",
    str(PROJECT / "results/11_manuscript_preparation/18B_second_round_manuscript_QA/18B_second_round_QA_revised_manuscript_v3.md"),
))
FIG_DIR = PROJECT / "results/11_manuscript_preparation/14C_figures_and_tables"
OUTPUT_DIR = Path(os.environ.get(
    "MANUSCRIPT_OUTPUT_DIR",
    str(PROJECT / "results/11_manuscript_preparation/18C_word_manuscript_with_figures"),
))
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DOCX = Path(os.environ.get(
    "MANUSCRIPT_OUTPUT_DOCX",
    str(OUTPUT_DIR / "18C_second_round_revised_manuscript_with_figures_v3.docx"),
))

FIGURES = {
    "## Core program support in paired sample-level validation": {
        "path": FIG_DIR / "14C_Figure2_module_support_by_source.png",
        "caption": "Figure 2 | Module-level support by source. Source-separated module support is shown for PRJNA607098 (12 paired sample-state units) and GSE130973 (5 subjects). PRJNA607098 bars use the strict sample-level support rule applied to the frozen candidate module; GSE130973 bars show the descriptive positive-subject fraction used in the external triangulation. The sources were not pooled, and the fractions are not equivalent confirmatory tests.",
    },
    "## Independent GSE130973 external triangulation": {
        "path": FIG_DIR / "14C_Figure3_candidate_support_cross_source.png",
        "caption": "Figure 3 | Candidate support across sources. Each tile represents one frozen candidate in one source. Colour indicates the source-specific support status. The PRJNA607098 column uses the strict 12-unit rule; the GSE130973 column uses the descriptive 4/5 directional rule. The figure displays concordance and discordance rather than an independent replication claim.",
    },
    "## Competition-robustness boundary": {
        "path": FIG_DIR / "14C_Figure4_competition_robustness_boundary.png",
        "caption": "Figure 4 | Competition robustness boundary. Bars show the competitor-adjusted residual median for the integrin/focal-adhesion core module in PRJNA607098 across four competing state programs (n = 12 paired sample-state units). Negative residual medians and failed 8/12 support do not prove that integrin biology is absent; they indicate that the current data do not establish specificity from the competing states.",
    },
}


def set_run_font(run, name="Calibri", size=11, color="000000", bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.widow_control = True


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color="6B7280")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_inline_markdown(paragraph, text, size=11, color="000000"):
    # Minimal inline Markdown: bold, italic, and code spans.
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, color=color, italic=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=size - 0.5, color="374151")
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, color=color)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph(style="Normal")
    set_paragraph_spacing(p, before=0, after=6, line=1.10)
    add_inline_markdown(p, text)
    return p


def add_figure(doc, figure_info):
    path = figure_info["path"]
    if not path.exists():
        raise FileNotFoundError(f"Figure not found: {path}")
    image_p = doc.add_paragraph()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.paragraph_format.keep_with_next = True
    set_paragraph_spacing(image_p, before=8, after=3, line=1.0)
    run = image_p.add_run()
    inline_shape = run.add_picture(str(path), width=Inches(6.25))
    alt_text = figure_info.get("alt", figure_info.get("caption", "Scientific figure"))
    inline_shape._inline.docPr.set("descr", alt_text)
    inline_shape._inline.docPr.set("title", alt_text.split(".", 1)[0][:250])

    cap = doc.add_paragraph(style="Figure Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(cap, before=3, after=10, line=1.0)
    add_inline_markdown(cap, figure_info["caption"], size=9, color="374151")


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10
        style.paragraph_format.keep_with_next = True

    if "Figure Caption" not in styles:
        caption_style = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption_style = styles["Figure Caption"]
    caption_style.base_style = styles["Normal"]
    caption_style.font.name = "Calibri"
    caption_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption_style.font.size = Pt(9)
    caption_style.font.color.rgb = RGBColor.from_string("374151")
    caption_style.paragraph_format.space_before = Pt(3)
    caption_style.paragraph_format.space_after = Pt(10)
    caption_style.paragraph_format.line_spacing = 1.0


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(header, before=0, after=0, line=1.0)
    header_text = os.environ.get(
        "MANUSCRIPT_HEADER",
        "Human fascia–muscle mechanosensitivity | Second-round revised manuscript",
    )
    hrun = header.add_run(header_text)
    set_run_font(hrun, size=8.5, color="6B7280")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(footer, before=0, after=0, line=1.0)
    add_page_number(footer)


def add_title_block(doc, title, status):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=8, after=8, line=1.0)
    run = p.add_run(title)
    set_run_font(run, size=20, color="0B2545", bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=18, line=1.0)
    run = p.add_run(status)
    set_run_font(run, size=10, color="6B7280", italic=True)


def build_doc():
    text = INPUT_MD.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    configure_styles(doc)
    configure_page(doc)

    title = lines[0].removeprefix("# ").strip()
    status_line = next((line for line in lines if line.startswith("**Manuscript status:**")), "")
    status = (
        re.sub(r"^\*\*Manuscript status:\*\*\s*", "", status_line).strip()
        or os.environ.get("MANUSCRIPT_STATUS", "Second-round revised manuscript")
    )
    add_title_block(doc, title, status)

    inserted_figures = set()
    current_heading = None
    skip_html = False
    for line in lines[1:]:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("<!--"):
            skip_html = not stripped.endswith("-->")
            continue
        if skip_html:
            if stripped.endswith("-->"):
                skip_html = False
            continue
        if stripped.startswith("**Manuscript status:**"):
            continue

        if stripped.startswith("### "):
            if current_heading in FIGURES and current_heading not in inserted_figures:
                add_figure(doc, FIGURES[current_heading])
                inserted_figures.add(current_heading)
            p = doc.add_paragraph(style="Heading 3")
            add_inline_markdown(p, stripped[4:], size=12, color="1F4D78")
            current_heading = stripped
            continue
        if stripped.startswith("## "):
            if current_heading in FIGURES and current_heading not in inserted_figures:
                add_figure(doc, FIGURES[current_heading])
                inserted_figures.add(current_heading)
            p = doc.add_paragraph(style="Heading 2")
            add_inline_markdown(p, stripped[3:], size=13, color="2E74B5",)
            current_heading = stripped
            continue
        if stripped.startswith("# "):
            continue
        if re.match(r"^\d+\.\s+", stripped):
            body = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            set_paragraph_spacing(p, before=0, after=4, line=1.10)
            add_inline_markdown(p, body)
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph_spacing(p, before=0, after=4, line=1.10)
            add_inline_markdown(p, stripped[2:])
            continue
        if stripped.startswith("> "):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Inches(0.25)
            set_paragraph_spacing(p, before=0, after=6, line=1.10)
            add_inline_markdown(p, stripped[2:], color="374151")
            continue
        if stripped == "." and doc.paragraphs:
            # Preserve punctuation from the Markdown source without creating an isolated line.
            doc.paragraphs[-1].add_run(".")
            continue
        if stripped.startswith("|"):
            # The manuscript contains no active data tables; preserve any table text as prose if encountered.
            add_body_paragraph(doc, stripped.strip("|"))
            continue

        add_body_paragraph(doc, stripped)

    if current_heading in FIGURES and current_heading not in inserted_figures:
        add_figure(doc, FIGURES[current_heading])
        inserted_figures.add(current_heading)

    # If a section had no body text, still insert its figure after the heading.
    for heading, info in FIGURES.items():
        if heading not in inserted_figures:
            raise RuntimeError(f"Figure insertion point not found: {heading}")

    doc.core_properties.title = title
    doc.core_properties.subject = "Second-round evidence-grounded manuscript draft"
    doc.core_properties.keywords = "fascia, mechanosensitivity, fibroblast, mechanotransduction"
    doc.core_properties.comments = "Generated from the QA-passed Markdown manuscript; figures inserted from the project figure package."
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    path = build_doc()
    print(path)
