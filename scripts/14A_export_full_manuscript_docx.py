"""Export the current 14A Markdown manuscript draft to a styled DOCX.

This is a formatting conversion only. Scientific claims and source wording are
preserved; the document is not a journal-specific submission file.
"""

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path(r".")
SOURCE = PROJECT / "results/11_manuscript_preparation/14A_full_manuscript_draft/14A_full_manuscript_draft_v1.md"
OUTPUT = PROJECT / "results/11_manuscript_preparation/14A_full_manuscript_draft/14A_full_manuscript_draft_v2_with_figures.docx"
FIGURE_DIR = PROJECT / "results/11_manuscript_preparation/14C_figures_and_tables"

FONT = "Calibri"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "666666"
INK = "0B2545"


FIGURE_SPECS = [
    {
        "number": "Figure 2",
        "anchor": "PIEZO2 was not supported in the PRJNA607098 paired analysis",
        "image": FIGURE_DIR / "14C_Figure2_module_support_by_source.png",
        "title": "Module-level support by source",
        "legend": (
            "Source-separated module support is shown for PRJNA607098 (12 paired "
            "sample-state units) and GSE130973 (5 subjects). PRJNA607098 bars use "
            "the strict sample-level support rule applied to the frozen candidate "
            "module; GSE130973 bars show the descriptive positive-subject fraction "
            "used in the external triangulation. Numerators and denominators are "
            "printed above bars. The sources were not pooled, and the fractions are "
            "not equivalent confirmatory tests."
        ),
    },
    {
        "number": "Figure 3",
        "anchor": "At the candidate level, integrin/focal-adhesion directional support",
        "image": FIGURE_DIR / "14C_Figure3_candidate_support_cross_source.png",
        "title": "Candidate support across sources",
        "legend": (
            "Each tile represents one frozen candidate in one source. Colour indicates "
            "the source-specific support status. Text reports the source-specific median "
            "directional difference and is not intended for cross-source magnitude "
            "comparison. The PRJNA607098 column uses the strict 12-unit rule; the "
            "GSE130973 column uses the descriptive 4/5 directional rule. The figure "
            "therefore displays concordance and discordance rather than an independent "
            "replication claim."
        ),
    },
    {
        "number": "Figure 4",
        "anchor": "The integrated interpretation was therefore restricted to a module-level",
        "image": FIGURE_DIR / "14C_Figure4_competition_robustness_boundary.png",
        "title": "Competition robustness boundary",
        "legend": (
            "Bars show the competitor-adjusted residual median for the integrin/focal- "
            "adhesion core module in PRJNA607098 across four competing state programs "
            "(n = 12 paired sample-state units). Text reports the number of positive "
            "residual units. Negative residual medians and failed 8/12 support do not "
            "prove that integrin biology is absent; they indicate that the current data "
            "do not establish specificity from the competing states."
        ),
    },
]


def set_run_font(run, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name=FONT, size=11, color="000000", bold=False, italic=False):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    style.font.italic = italic
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size=11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    h1 = styles["Heading 1"]
    set_style_font(h1, size=16, color=BLUE, bold=True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.keep_together = True

    h2 = styles["Heading 2"]
    set_style_font(h2, size=13, color=BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.keep_together = True

    h3 = styles["Heading 3"]
    set_style_font(h3, size=12, color=DARK_BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True
    h3.paragraph_format.keep_together = True

    title_style = styles.add_style("Manuscript Title", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(title_style, size=22, color=INK, bold=True)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(6)
    title_style.paragraph_format.keep_with_next = True

    status_style = styles.add_style("Draft Status", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(status_style, size=10, color=MUTED, italic=True)
    status_style.paragraph_format.space_before = Pt(0)
    status_style.paragraph_format.space_after = Pt(16)
    status_style.paragraph_format.line_spacing = 1.15

    keyword_style = styles.add_style("Keywords", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(keyword_style, size=10, color=MUTED)
    keyword_style.paragraph_format.space_before = Pt(2)
    keyword_style.paragraph_format.space_after = Pt(14)
    keyword_style.paragraph_format.line_spacing = 1.15

    figure_caption_style = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(figure_caption_style, size=9.5, color="222222")
    figure_caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    figure_caption_style.paragraph_format.space_before = Pt(2)
    figure_caption_style.paragraph_format.space_after = Pt(14)
    figure_caption_style.paragraph_format.line_spacing = 1.15
    figure_caption_style.paragraph_format.keep_together = True

    code_style = styles.add_style("Inline Code", WD_STYLE_TYPE.CHARACTER)
    set_style_font(code_style, name="Courier New", size=9, color="444444")

    # Quiet running label and centered page number; no decorative rule.
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("Human fascia–muscle mechanosensitivity study | Initial manuscript draft")
    set_run_font(hr, size=8.5, color=MUTED)
    hp.paragraph_format.space_after = Pt(0)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Page ")
    set_run_font(fr, size=9, color=MUTED)
    add_page_field(fp)
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)


def add_inline_markdown(paragraph, text):
    """Add modest inline Markdown formatting without exposing Markdown syntax."""
    token_re = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Courier New", size=9, color="444444")
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, italic=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run)


def add_body_paragraph(doc, text, style="Normal"):
    paragraph = doc.add_paragraph(style=style)
    add_inline_markdown(paragraph, text)
    # Long inline paths can create excessive word spacing under justified text.
    if style == "Normal" and ("D:/" in text or "D:\\" in text):
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return paragraph


def add_figure_block(doc, spec):
    """Insert a full-width figure followed by its evidence-bounded legend."""
    if not spec["image"].exists():
        raise FileNotFoundError(f"Figure image not found: {spec['image']}")

    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.space_before = Pt(4)
    image_paragraph.paragraph_format.space_after = Pt(4)
    image_paragraph.paragraph_format.keep_with_next = True
    image_run = image_paragraph.add_run()
    image_run.add_picture(str(spec["image"]), width=Inches(6.5))

    caption = doc.add_paragraph(style="Figure Caption")
    title_run = caption.add_run(f"{spec['number']}. {spec['title']}. ")
    set_run_font(title_run, size=9.5, color="222222", bold=True)
    legend_run = caption.add_run(spec["legend"])
    set_run_font(legend_run, size=9.5, color="222222")


def clean_source_line(text):
    # Preserve the scientific text while removing only Markdown syntax markers.
    return text.strip()


def export_docx():
    if not SOURCE.exists():
        raise FileNotFoundError(f"Source Markdown not found: {SOURCE}")

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "Human fascia–muscle mechanosensitivity: initial manuscript draft"
    doc.core_properties.subject = "Evidence-grounded manuscript draft"
    doc.core_properties.author = ""
    doc.core_properties.comments = "Formatting conversion from the project Markdown draft; not journal-specific formatting."

    paragraph_buffer = []
    first_title_done = False
    last_heading = None
    inserted_figures = set()

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            text = " ".join(x.strip() for x in paragraph_buffer).strip()
            if text:
                style = "Keywords" if text.startswith("Keywords:") else "Normal"
                add_body_paragraph(doc, text, style=style)
                for spec in FIGURE_SPECS:
                    if spec["number"] not in inserted_figures and spec["anchor"] in text:
                        add_figure_block(doc, spec)
                        inserted_figures.add(spec["number"])
            paragraph_buffer = []

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            flush_paragraph()
            continue

        title_match = re.match(r"^#\s+(.+)$", line)
        if title_match and not first_title_done:
            flush_paragraph()
            title = title_match.group(1).strip()
            p = doc.add_paragraph(style="Manuscript Title")
            run = p.add_run(title)
            set_run_font(run, size=22, color=INK, bold=True)
            first_title_done = True
            continue

        # Two section labels lost their Markdown heading prefix during the
        # manuscript assembly step. Restore them as real Word subheadings.
        artifact_heading = {
            "n and analytical scope": "Design and analytical scope",
            "nance and validation structure": "Provenance and validation structure",
        }.get(line.strip())
        if artifact_heading is not None:
            flush_paragraph()
            add_body_paragraph(doc, artifact_heading, style="Heading 2")
            last_heading = artifact_heading
            continue

        # Section/subsection hierarchy from the Markdown source.
        heading_match = re.match(r"^(#{2,6})\s+(.+)$", line)
        if heading_match:
            flush_paragraph()
            level = min(len(heading_match.group(1)) - 1, 3)
            text = heading_match.group(2).strip()
            # These are assembly fragments in the current Markdown draft.
            text = {
                "n and analytical scope": "Design and analytical scope",
                "nance and validation structure": "Provenance and validation structure",
            }.get(text, text)
            # Keep both source conclusion paragraphs under one heading.
            if text == "Conclusion" and last_heading == "Conclusion":
                continue
            add_body_paragraph(doc, text, style=f"Heading {level}")
            last_heading = text
            continue

        # The manuscript status line is kept as a distinct metadata paragraph.
        status_match = re.match(r"^\*\*Manuscript status:\*\*\s*(.+)$", line)
        if status_match:
            flush_paragraph()
            p = doc.add_paragraph(style="Draft Status")
            label = p.add_run("Manuscript status: ")
            set_run_font(label, size=10, color=MUTED, bold=True)
            value = p.add_run(status_match.group(1).strip())
            set_run_font(value, size=10, color=MUTED, italic=True)
            continue

        # An isolated "n" is a known section-assembly fragment after Discussion.
        if line.strip() == "n":
            continue

        list_match = re.match(r"^\s*([-*+]|\d+\.)\s+(.+)$", line)
        if list_match:
            flush_paragraph()
            marker = list_match.group(1)
            text = list_match.group(2).strip()
            style = "List Number" if marker.endswith(".") and marker[:-1].isdigit() else "List Bullet"
            add_body_paragraph(doc, text, style=style)
            continue

        # Preserve all other lines as paragraph content; source line wrapping is
        # normalized only within the paragraph, not across blank-line boundaries.
        paragraph_buffer.append(clean_source_line(line))

    flush_paragraph()

    missing_figures = [spec["number"] for spec in FIGURE_SPECS if spec["number"] not in inserted_figures]
    if missing_figures:
        raise RuntimeError(f"Figure insertion anchors were not found: {', '.join(missing_figures)}")

    # Update fields on open where supported by Word/LibreOffice.
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Created DOCX: {OUTPUT}")


if __name__ == "__main__":
    export_docx()
