"""Build the BMC Genomics-formatted Word manuscript with figures.

The source Markdown is prepared by 23A_prepare_bmc_genomics_manuscript.py.
The output is a single-column, double-spaced, line-numbered editable DOCX.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


PROJECT = Path(r".")
INPUT_MD = PROJECT / "results/11_manuscript_preparation/23A_BMC_Genomics_format/23A_BMC_Genomics_manuscript_v1.md"
OUTPUT_DIR = PROJECT / "results/11_manuscript_preparation/23B_BMC_Genomics_word_manuscript"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DOCX = OUTPUT_DIR / "23B_BMC_Genomics_manuscript_v2_clean_headings_with_figures.docx"


FIGURES = {
    "### Provenance and primary validation structure": {
        "path": PROJECT / "results/11_manuscript_preparation/20B_reviewer_response_figures/20B_Figure1_evidence_workflow.png",
        "caption": "Figure 1. Evidence workflow and inferential units. The staged design separates discovery, paired atlas validation, independent subject-level triangulation, direct mechanical perturbation, regulatory/stiffness context and bounded synthesis. Each source retains its own biological or inferential denominator.",
    },
    "### Core modules were consistently positive in PRJNA607098": {
        "path": PROJECT / "results/11_manuscript_preparation/20B_reviewer_response_figures/20B_Figure2_sample_subject_module_contrasts.png",
        "caption": "Figure 2. Primary-module contrasts at the correct inferential unit. PRJNA607098 points are 12 reconstructed SRS sample units comparing F7 with pooled non-F7 cells within sample. GSE130973 points are five subjects comparing marker-defined candidate fibroblast clusters with other clusters within subject. Black summaries are medians and IQRs; source-specific free y-axis scales are intentional.",
    },
    "### Competition analysis limited integrin specificity": {
        "path": PROJECT / "results/11_manuscript_preparation/20B_reviewer_response_figures/20B_Figure4_competition_residuals.png",
        "caption": "Figure 3. Competition residuals expose the specificity boundary. Within PRJNA607098, each core-module contrast was regressed on each competitor-module contrast. The prespecified primary gate requires at least 8 of 12 residuals above zero; a secondary sensitivity summary requires at least 9 of 12. Integrin versus ECM remodelling and three actomyosin comparisons meet only the minimum 8/12 boundary.",
    },
    "### Independent skin data supported the modules but not candidate uniformity": {
        "path": PROJECT / "results/11_manuscript_preparation/20B_reviewer_response_figures/20B_Figure3_frozen_candidate_raw_contrasts.png",
        "caption": "Figure 4. Frozen candidate effects are source dependent. All frozen candidates are shown without post hoc removal. Points are sample- or subject-level target-minus-comparator contrasts; black diamonds and bars show median and IQR. Source separation preserves distinct expression scales and denominators.",
    },
    "### Direct mechanical tension supported actomyosin/Rho but also cell cycle": {
        "path": PROJECT / "results/11_manuscript_preparation/20B_reviewer_response_figures/20B_Figure5_GSE300230_tension_response.png",
        "caption": "Figure 5. Direct mechanical tension activates actomyosin and cell-cycle programmes. Model-estimated module-score contrasts compare tension with relaxed collagen without TGF-beta in each GSE300230 cell line. Labels give CAMERA BH q values. The concurrent cell-cycle response is an interpretation boundary; the two cell lines are not independent donors or age replicates.",
    },
    "### Module-level support was more reproducible than candidate-gene support": {
        "path": PROJECT / "results/11_manuscript_preparation/21B_evidence_hierarchy_figure/21B_Figure6_evidence_hierarchy.png",
        "caption": "Figure 6. Module-level support is more reproducible than individual-gene support across sources. The matrix summarises the frozen evidence hierarchy across primary paired validation, independent skin triangulation, direct tension, TEAD-linked regulatory cross-validation and cross-tissue stiffness analysis. Partial or discordant cells retain co-active competitor programmes, limited replication and source-dependent candidates. Colours are descriptive evidence classes, not a universal grading scale.",
    },
}


def set_run_font(run, name="Times New Roman", size=12, color="000000", bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=0, line=2.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.widow_control = True


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    set_run_font(run, size=10, color="666666")


def enable_line_numbers(section):
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:lnNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    line_numbers = OxmlElement("w:lnNumType")
    line_numbers.set(qn("w:countBy"), "1")
    line_numbers.set(qn("w:start"), "1")
    line_numbers.set(qn("w:distance"), "360")
    line_numbers.set(qn("w:restart"), "continuous")
    sect_pr.append(line_numbers)


def add_inline_markdown(paragraph, text, size=12, color="000000"):
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, color=color, italic=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=max(size - 1, 9), color="333333")
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, color=color)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 2.0

    for name, size, before, after in [
        ("Heading 1", 14, 12, 4),
        ("Heading 2", 12, 8, 2),
        ("Heading 3", 12, 6, 0),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        # Do not set keep-with-next here: Word displays a small black square in
        # the left margin for this pagination property when formatting marks
        # are enabled. The manuscript remains readable without that marker.
        style.paragraph_format.keep_with_next = False

    if "BMC Figure Caption" not in styles:
        caption_style = styles.add_style("BMC Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption_style = styles["BMC Figure Caption"]
    caption_style.base_style = styles["Normal"]
    caption_style.font.name = "Times New Roman"
    caption_style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    caption_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    caption_style.font.size = Pt(10)
    caption_style.font.color.rgb = RGBColor.from_string("333333")
    caption_style.paragraph_format.space_before = Pt(2)
    caption_style.paragraph_format.space_after = Pt(8)
    caption_style.paragraph_format.line_spacing = 1.0


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.header_distance = Mm(12)
    section.footer_distance = Mm(12)
    enable_line_numbers(section)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(header, line=1.0)
    run = header.add_run("BMC Genomics | Research Article")
    set_run_font(run, size=9, color="666666")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(footer, line=1.0)
    add_page_number(footer)


def add_title_page(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=4, after=6, line=1.0)
    run = p.add_run(title)
    set_run_font(run, size=16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=8, line=1.0)
    run = p.add_run("Research Article")
    set_run_font(run, size=11, italic=True, color="444444")

    for text in [
        "Authors: [AUTHOR 1 FULL NAME]1, [AUTHOR 2 FULL NAME]2",
        "1. [Department, Institution, City, Country]",
        "2. [Department, Institution, City, Country]",
        "Corresponding author: [FULL NAME, EMAIL ADDRESS]",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, line=1.0, after=2)
        run = p.add_run(text)
        set_run_font(run, size=11, color="333333")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=6, after=12, line=1.0)
    run = p.add_run("Author details and declarations marked for completion must be finalised before submission.")
    set_run_font(run, size=9, color="666666", italic=True)


def add_body_paragraph(doc, text, size=12, after=0):
    p = doc.add_paragraph(style="Normal")
    set_paragraph_spacing(p, after=after, line=2.0)
    add_inline_markdown(p, text, size=size)
    return p


def add_reference_paragraph(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Mm(6)
    p.paragraph_format.first_line_indent = Mm(-6)
    set_paragraph_spacing(p, line=2.0)
    add_inline_markdown(p, text, size=11)
    return p


def add_figure(doc, info):
    path = info["path"]
    if not path.exists():
        raise FileNotFoundError(f"Figure not found: {path}")
    image_p = doc.add_paragraph()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.paragraph_format.keep_with_next = True
    set_paragraph_spacing(image_p, before=4, after=2, line=1.0)
    run = image_p.add_run()
    shape = run.add_picture(str(path), width=Inches(6.45))
    alt = info["caption"]
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("title", alt.split(".", 1)[0][:250])

    caption = doc.add_paragraph(style="BMC Figure Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_inline_markdown(caption, info["caption"], size=10, color="333333")


def build_doc():
    text = INPUT_MD.read_text(encoding="utf-8")
    lines = text.splitlines()
    title = lines[0].removeprefix("# ").strip()
    doc = Document()
    # python-docx starts with one empty body paragraph. Removing it prevents a
    # blank first line and makes continuous line numbering begin at the title.
    if doc.paragraphs and not doc.paragraphs[0].text:
        empty = doc.paragraphs[0]._element
        empty.getparent().remove(empty)
    configure_styles(doc)
    configure_page(doc)
    add_title_page(doc, title)

    inserted = set()
    current_heading = None
    in_references = False
    started_main = False

    def flush_figure():
        if current_heading in FIGURES and current_heading not in inserted:
            add_figure(doc, FIGURES[current_heading])
            inserted.add(current_heading)

    for raw_line in lines[1:]:
        stripped = raw_line.strip()
        if not stripped:
            continue
        if stripped.startswith("<!--"):
            continue
        if stripped.startswith("!["):
            continue
        if stripped.startswith("**Figure ") and stripped.endswith("**"):
            continue
        if not started_main and re.match(r"^\d+\.\s+\[", stripped):
            continue
        if stripped.startswith("## ") or stripped.startswith("### "):
            flush_figure()
            started_main = True
            if stripped == "## References":
                in_references = True
            current_heading = stripped
            level = 1 if stripped.startswith("## ") else 2
            style = "Heading 1" if level == 1 else "Heading 2"
            p = doc.add_paragraph(style=style)
            set_paragraph_spacing(p, line=1.0)
            text_heading = stripped[3:] if level == 1 else stripped[4:]
            add_inline_markdown(p, text_heading, size=14 if level == 1 else 12)
            continue
        if stripped.startswith("# "):
            continue
        if stripped.startswith("*Research Article*") or stripped.startswith("**Authors:**") or stripped.startswith("**Corresponding author:**"):
            continue
        if re.match(r"^\d+\.\s+\[", stripped) and not in_references:
            add_body_paragraph(doc, stripped)
            continue
        if in_references and re.match(r"^\d+\.\s+", stripped):
            add_reference_paragraph(doc, stripped)
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph_spacing(p, line=2.0)
            add_inline_markdown(p, stripped[2:])
            continue
        if stripped.startswith("> "):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Mm(6)
            set_paragraph_spacing(p, line=2.0)
            add_inline_markdown(p, stripped[2:], color="444444")
            continue
        if started_main:
            add_body_paragraph(doc, stripped)

    flush_figure()
    missing = [heading for heading in FIGURES if heading not in inserted]
    if missing:
        raise RuntimeError(f"Figure insertion point(s) not found: {missing}")

    doc.core_properties.title = title
    doc.core_properties.subject = "BMC Genomics Research article manuscript"
    doc.core_properties.keywords = "fibroblast, fascia, mechanotransduction, genomics"
    doc.core_properties.comments = "BMC Genomics-formatted manuscript generated from the QA-passed project draft."
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    print(build_doc())
