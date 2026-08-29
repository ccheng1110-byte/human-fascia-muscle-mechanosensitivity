from __future__ import annotations

import csv
import hashlib
import re
import zipfile
from pathlib import Path

from docx import Document
from pypdf import PdfReader


PROJECT_DIR = Path(r".")
INPUT_DIR = PROJECT_DIR / "data" / "metadata" / "GSE273293" / "supporting_information"
OUTPUT_DIR = (
    PROJECT_DIR
    / "results"
    / "07_deep_fascia_recovery"
    / "GSE273293"
    / "09B_processed_data_recovery_audit"
    / "supporting_information_extraction"
)
RENDER_DIR = (
    PROJECT_DIR
    / "results"
    / "07_deep_fascia_recovery"
    / "GSE273293"
    / "09B_processed_data_recovery_audit"
    / "supporting_information_render"
)

SOURCE_URLS = {
    "12967_2024_5889_MOESM1_ESM.docx": (
        "https://static-content.springer.com/esm/"
        "art%3A10.1186%2Fs12967-024-05889-y/MediaObjects/"
        "12967_2024_5889_MOESM1_ESM.docx"
    ),
    "12967_2024_5889_MOESM2_ESM.docx": (
        "https://static-content.springer.com/esm/"
        "art%3A10.1186%2Fs12967-024-05889-y/MediaObjects/"
        "12967_2024_5889_MOESM2_ESM.docx"
    ),
}

KEY_PATTERN = re.compile(
    r"(?:GSE273293|PRJNA1206333|GMC[_ -]?\d+|Con[_ -]?\d+|control[_ -]?\d+|"
    r"barcode|sample(?:_id| id)?|Seurat|\.rds\b|\.h5ad\b|matrix|GitHub|repository|code)",
    re.IGNORECASE,
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest().upper()


def iter_block_text(document: Document):
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            yield "paragraph", str(index), text

    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            text = "\t".join(cells).strip()
            if text:
                yield "table", f"{table_index}.{row_index}", text

    for section_index, section in enumerate(document.sections, start=1):
        for area_name, area in (("header", section.header), ("footer", section.footer)):
            for paragraph_index, paragraph in enumerate(area.paragraphs, start=1):
                text = paragraph.text.strip()
                if text:
                    yield area_name, f"{section_index}.{paragraph_index}", text


def inspect_docx(path: Path):
    document = Document(path)
    blocks = list(iter_block_text(document))

    text_path = OUTPUT_DIR / f"{path.stem}_extracted_text.txt"
    with text_path.open("w", encoding="utf-8", newline="\n") as handle:
        for kind, location, text in blocks:
            handle.write(f"[{kind} {location}] {text}\n")

    keyword_rows = []
    for kind, location, text in blocks:
        matches = sorted({match.group(0) for match in KEY_PATTERN.finditer(text)}, key=str.lower)
        if matches:
            keyword_rows.append(
                {
                    "document": path.name,
                    "block_type": kind,
                    "location": location,
                    "matched_terms": "; ".join(matches),
                    "text": text,
                }
            )

    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        embedded = [
            name
            for name in names
            if name.startswith("word/embeddings/")
            or name.startswith("word/media/")
            or name.lower().endswith((".csv", ".tsv", ".xlsx", ".xls", ".rds", ".h5ad"))
        ]

        relationship_rows = []
        for name in names:
            if not name.endswith(".rels"):
                continue
            xml = archive.read(name).decode("utf-8", errors="replace")
            for match in re.finditer(r'Target="([^"]+)"', xml):
                target = match.group(1)
                if target.startswith(("http://", "https://", "ftp://", "mailto:")):
                    relationship_rows.append(
                        {"document": path.name, "relationship_part": name, "target": target}
                    )

    return {
        "manifest": {
            "file_name": path.name,
            "source_url": SOURCE_URLS[path.name],
            "bytes": path.stat().st_size,
            "sha256": sha256(path),
            "docx_valid": True,
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
            "keyword_hit_blocks": len(keyword_rows),
            "external_relationships": len(relationship_rows),
            "embedded_or_media_parts": len(embedded),
        },
        "keywords": keyword_rows,
        "relationships": relationship_rows,
        "embedded": [
            {"document": path.name, "archive_part": item}
            for item in embedded
        ],
    }


def write_csv(path: Path, rows: list[dict], fieldnames: list[str]):
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    results = []
    for name in SOURCE_URLS:
        path = INPUT_DIR / name
        if not path.is_file():
            raise FileNotFoundError(path)
        results.append(inspect_docx(path))

    write_csv(
        INPUT_DIR / "GSE273293_supporting_information_manifest_v1.csv",
        [result["manifest"] for result in results],
        list(results[0]["manifest"].keys()),
    )
    write_csv(
        OUTPUT_DIR / "GSE273293_supporting_information_keyword_hits_v1.csv",
        [row for result in results for row in result["keywords"]],
        ["document", "block_type", "location", "matched_terms", "text"],
    )
    write_csv(
        OUTPUT_DIR / "GSE273293_supporting_information_external_links_v1.csv",
        [row for result in results for row in result["relationships"]],
        ["document", "relationship_part", "target"],
    )
    write_csv(
        OUTPUT_DIR / "GSE273293_supporting_information_embedded_parts_v1.csv",
        [row for result in results for row in result["embedded"]],
        ["document", "archive_part"],
    )

    for document_number in (1, 2):
        stem = f"12967_2024_5889_MOESM{document_number}_ESM"
        pdf_path = RENDER_DIR / f"MOESM{document_number}" / f"{stem}.pdf"
        if not pdf_path.is_file():
            continue
        reader = PdfReader(pdf_path)
        pdf_text_path = OUTPUT_DIR / f"{stem}_rendered_pdf_text.txt"
        with pdf_text_path.open("w", encoding="utf-8", newline="\n") as handle:
            for page_number, page in enumerate(reader.pages, start=1):
                handle.write(f"===== PAGE {page_number} =====\n")
                handle.write(page.extract_text() or "")
                handle.write("\n")

    for result in results:
        print(result["manifest"])
    print(f"Extraction output: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
